"""Generate DOCX: Mapeamento Completo de Fluxos Financeiros — MAEZO Healthcare Platform."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_shading(cell, "EDF2F7")

    # Column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    return table


def add_gap_box(doc, text):
    """Add a highlighted GAP notice."""
    p = doc.add_paragraph()
    run = p.add_run("⚠ GAP: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x66, 0x33, 0x00)


# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Mapeamento Completo de\nFluxos Financeiros")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("MAEZO Healthcare Platform")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x7A, 0xB5)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    "Levantamento de todos os processos BPMN, workers, DMN e serviços\n"
    "relacionados a pagamento, cobrança, nota fiscal, conciliação,\n"
    "cancelamento e operações financeiras.\n\n"
    "Contexto: Portal de Cobrança Automática — Exames de Imagem"
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
datep = doc.add_paragraph()
datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = datep.add_run(f"Data: {date.today().strftime('%d/%m/%Y')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Sumário", level=1)
toc_items = [
    "1. Visão Geral",
    "2. Cobrança / Billing — Geração de Cobranças e Faturamento TISS",
    "3. Pagamento / Payment — Self-Service do Paciente",
    "4. Nota Fiscal / Invoice — Emissão de NFS-e",
    "5. Conciliação Bancária / Reconciliation",
    "6. Cancelamento / Refund / Estorno",
    "7. Cobrança de Inadimplentes / Dunning",
    "8. Glosa / Denial Management",
    "9. Maximização e Revenue Optimization",
    "10. Financial Clearance (Patient Access)",
    "11. Produção Clínica e Precificação",
    "12. Resumo: O Que Existe vs. Gaps para o Portal",
    "13. Recomendações de Implementação",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. VISAO GERAL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Visão Geral", level=1)
doc.add_paragraph(
    "O ciclo de receita (Revenue Cycle) da plataforma MAEZO abrange 14+ processos BPMN, "
    "~150 workers, 150+ tabelas de decisão DMN e 12+ serviços de suporte. "
    "Este documento mapeia todos os componentes relacionados a operações financeiras, "
    "organizados em categorias funcionais."
)

doc.add_heading("Processo Orquestrador", level=2)
doc.add_paragraph(
    "SP-RC-000_Revenue_Cycle_Main.bpmn — Orquestra todos os subprocessos (SP-RC-001 a SP-RC-014) "
    "via call activities, com message events: encounter_created, clinical_doc_ready, "
    "discharge_completed, payment_reconciled, glosa_detected."
)

doc.add_heading("Domínios Envolvidos", level=2)
add_table(doc,
    ["Domínio", "Subdomínios", "Relevância Financeira"],
    [
        ["revenue_cycle", "billing, coding, collection, glosa, production", "Principal — todo o ciclo financeiro"],
        ["patient_access", "scheduling, registration, engagement", "Financial Clearance, pré-autorização"],
        ["platform_services", "analytics, integration, revenue_optimization", "Otimização, reconciliação, métricas"],
        ["contract_extraction", "FastAPI service", "Extração de regras contratuais"],
    ],
    col_widths=[4, 7, 7],
)

# ═══════════════════════════════════════════════════════════════════
# 2. COBRANCA / BILLING
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("2. Cobrança / Billing — Geração de Cobranças e Faturamento TISS", level=1)

doc.add_heading("2.1 Processos BPMN", level=2)
add_table(doc,
    ["Processo", "ID", "Descrição"],
    [
        ["SP-RC-006\nBilling Submission",
         "SP_RC_006",
         "Pipeline completo: validar claim → calcular cobranças → regras de contrato → "
         "gerar TISS XML → validar schema → submeter à operadora. Timer SLA 24h."],
        ["SP-RC-004\nClinical Production",
         "SP_RC_004",
         "Captura de procedimentos: capturar → enriquecer → validar compatibilidade → "
         "calcular quantidade → atribuir preços → checar autorização → persistir."],
        ["SP-RC-005\nCoding Audit",
         "SP_RC_005",
         "Codificação médica: extrair dados clínicos → sugerir TUSS/CID-10 → "
         "aplicar regras → detectar fraude → auditar → finalizar."],
        ["SP-RC-002\nPre-Service",
         "SP_RC_002",
         "Pré-atendimento: checar autorização → solicitar autorização → "
         "validar procedimento. Timer 48h para autorização."],
    ],
    col_widths=[4, 3, 11],
)

doc.add_heading("2.2 Workers de Billing (13 workers)", level=2)
doc.add_paragraph("Path: revenue_cycle/billing/workers/")
add_table(doc,
    ["Worker", "TOPIC", "Descrição"],
    [
        ["validate_claim", "billing.validate_claim", "Valida dados do claim (encounter, paciente, procedimentos)"],
        ["calculate_charges", "billing.calculate_charges", "Calcula valores das cobranças"],
        ["apply_contract_rules", "billing.apply_contract_rules", "Aplica regras contratuais (descontos, taxas)"],
        ["apply_discounts", "billing.apply_discounts", "Aplica descontos contratuais"],
        ["consolidate_charges", "billing.consolidate_charges", "Consolida linhas de cobrança em grupos"],
        ["group_by_guide", "billing.group_by_guide", "Agrupa procedimentos por tipo de guia TISS"],
        ["generate_tiss_xml", "billing.generate_tiss_xml", "Gera XML TISS para envio à operadora"],
        ["validate_tiss_schema", "billing.validate_tiss_schema", "Valida TISS XML contra schema XSD"],
        ["submit_to_payer", "billing.submit_to_payer", "Submete TISS XML à operadora"],
        ["handle_acknowledgment", "billing.handle_acknowledgment", "Processa ACK da operadora"],
        ["track_protocol", "billing.track_protocol", "Rastreia números de protocolo"],
        ["notify_submission_status", "billing.notify_submission_status", "Envia notificações de status"],
        ["retry_failed_submission", "billing.retry_failed_submission", "Retenta envios com backoff exponencial"],
    ],
    col_widths=[4.5, 5, 8.5],
)

doc.add_heading("2.3 Serviços de Billing", level=2)
add_table(doc,
    ["Serviço", "Classe", "Descrição"],
    [
        ["claim_submission_service.py", "ClaimSubmissionService", "Orquestra envio TISS e retry"],
        ["tiss_generation_service.py", "TISSGenerationService", "Gera XML TISS a partir de dados do claim"],
        ["tiss_validation_service.py", "TISSValidationService", "Valida XML TISS contra schema"],
        ["billing_rules_service.py", "BillingRulesService", "Decisões DMN agregadas (quantity, modifier, upcode)"],
        ["pricing_service.py", "PricingService", "Precificação contratual via federação DMN"],
    ],
    col_widths=[5.5, 5, 7.5],
)

doc.add_heading("2.4 Tabelas de Decisão DMN (80+ arquivos)", level=2)
doc.add_paragraph("Path: revenue_cycle/dmn/billing/")
add_table(doc,
    ["Subdiretório", "Qtd", "Propósito"],
    [
        ["quantity/", "5", "Validação de quantidade (bill_quantity_001..005)"],
        ["modifier/", "5", "Regras de modificadores (bill_modifier_001..005)"],
        ["bundle/", "5", "Precificação de pacotes (bill_bundle_001..005)"],
        ["bundle-ext/", "5", "Pacotes estendidos (bill_bundle_ext_001..005)"],
        ["taxa/", "5", "Cálculo de taxas (bill_taxa_001..005)"],
        ["opme/", "10", "Precificação OPME — implantes/materiais (bill_opme_001..010)"],
        ["material/", "7", "Regras de materiais (bill_material_001..007)"],
        ["med/", "3", "Faturamento de medicamentos (bill_med_001..003)"],
        ["diaria/", "5", "Regras de diárias (bill_diaria_001..005)"],
        ["time/", "5", "Faturamento por tempo (bill_time_001..005)"],
        ["specialty/", "2", "Faturamento por especialidade (bill_specialty_001..002)"],
        ["upcode/", "5", "Detecção de upcoding (bill_upcode_001..005)"],
        ["charges/", "1", "Regras de cálculo de cobrança"],
        ["validation/", "1", "Validação de claims"],
        ["submission/", "1", "Validação de submissão"],
        ["contract_rules/", "1", "Validação de regras contratuais"],
        ["discounts/", "1", "Aplicabilidade de descontos"],
        ["consolidation/", "1", "Regras de consolidação"],
        ["grouping/", "1", "Classificação de tipo de guia TISS"],
        ["tiss_generation/", "1", "Validação de geração TISS XML"],
        ["tiss_validation/", "1", "Validação de schema TISS"],
        ["notifications/", "1", "Roteamento de notificações"],
        ["retry/", "1", "Lógica de retry"],
        ["acknowledgment/", "1", "Processamento de ACK"],
        ["tracking/", "1", "Validação de rastreamento de protocolo"],
        ["federated/", "1", "Faturamento federado"],
    ],
    col_widths=[4, 1.5, 12.5],
)

# ═══════════════════════════════════════════════════════════════════
# 3. PAGAMENTO / PAYMENT
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("3. Pagamento / Payment — Self-Service do Paciente", level=1)

doc.add_heading("3.1 Processo BPMN Principal", level=2)
doc.add_paragraph(
    "SP-RC-014_Patient_Financial_SelfService.bpmn — FUNDAMENTO DO PORTAL"
)
doc.add_paragraph(
    "Este é o processo central para o portal de cobrança automática. "
    "Define 3 caminhos:"
)

add_table(doc,
    ["Path", "Nome", "Fluxo"],
    [
        ["A", "Pre-Visit Estimate",
         "Calcular copay → enviar estimativa → ACK do paciente (timeout 3 dias)"],
        ["B", "Bill Ready",
         "Gerar detalhes da conta → notificar paciente → ação: VIEW / PAY / QUESTION.\n"
         "PAY → subprocesso: gateway método (PAY_NOW → processar pagamento / "
         "PAY_LATER → plano de pagamento) → confirmação → recibo"],
        ["C", "Insurance Update",
         "Solicitar info do convênio → verificar cobertura → atualizar registro"],
    ],
    col_widths=[1.5, 4, 12.5],
)

doc.add_heading("3.2 Topics do Self-Service", level=2)
add_table(doc,
    ["TOPIC", "Descrição"],
    [
        ["financial.copay_estimate", "Cálculo de estimativa de copay"],
        ["financial.send_estimate", "Envio da estimativa ao paciente"],
        ["financial.bill_ready", "Preparação da conta"],
        ["financial.bill_notification", "Notificação de conta pronta"],
        ["financial.bill_detail", "Detalhamento da conta"],
        ["financial.billing_support", "Suporte a dúvidas de faturamento"],
        ["financial.payment_process", "Processamento do pagamento"],
        ["financial.payment_plan", "Criação de plano de pagamento"],
        ["financial.payment_confirmed", "Confirmação de pagamento"],
        ["financial.send_receipt", "Envio de recibo"],
        ["financial.insurance_request", "Solicitação de dados do convênio"],
        ["financial.verify_coverage", "Verificação de cobertura"],
        ["financial.update_record", "Atualização de registro financeiro"],
    ],
    col_widths=[6, 12],
)

doc.add_heading("3.3 DMN de Cash Operations (15 arquivos)", level=2)
doc.add_paragraph("Path: revenue_cycle/dmn/cash_operations/")
add_table(doc,
    ["Subdiretório", "Qtd", "Propósito"],
    [
        ["discount/", "3", "Regras de desconto à vista (cash_discount_001..003)"],
        ["estimate/", "3", "Regras de estimativa de custo (cash_estimate_001..003)"],
        ["estimates/", "4", "Regras de copay + copay_validation_adjudication"],
        ["payment/", "3", "Regras de método de pagamento (cash_payment_001..003)"],
        ["notifications/", "2", "Roteamento de notificações de autorização"],
    ],
    col_widths=[4, 1.5, 12.5],
)

doc.add_heading("3.4 Self-Service Eligibility DMN", level=2)
doc.add_paragraph(
    "selfservice_eligibility_rules.dmn (platform_services/dmn/communication/) — "
    "Determina elegibilidade para self-service por tipo de ação (confirm, cancel, "
    "reschedule, pay, question), nível de risco do paciente, janela de tempo e complexidade."
)

doc.add_heading("3.5 Enums de Pagamento", level=2)
doc.add_paragraph("Path: revenue_cycle/collection/enums.py")
add_table(doc,
    ["Enum", "Valores"],
    [
        ["PaymentStatus", "pending → received → validated → allocated → partially_allocated → reconciled → rejected → reversed"],
        ["PaymentType", "full, partial, advance, refund, adjustment"],
        ["PaymentMethod", "bank_transfer, boleto, pix, deposit, check, credit_card"],
        ["AgingBucket", "current, 0-30, 31-60, 61-90, 91-120, 121-180, 180+"],
        ["ReconciliationPeriod", "daily, weekly, monthly"],
    ],
    col_widths=[5, 13],
)

# ═══════════════════════════════════════════════════════════════════
# 4. NOTA FISCAL
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("4. Nota Fiscal / Invoice — Emissão de NFS-e", level=1)

doc.add_heading("4.1 O Que Existe", level=2)
doc.add_paragraph(
    "Não existe um BPMN dedicado para emissão de NFS-e (Nota Fiscal de Serviço Eletrônica). "
    "O que existe é a geração de TISS XML, que funciona como a \"nota\" para operadoras de saúde:"
)
add_table(doc,
    ["Componente", "Arquivo", "Descrição"],
    [
        ["Worker", "generate_tiss_xml_worker_v2.py", "Gera TISS XML (TOPIC: billing.generate_tiss_xml)"],
        ["Service", "tiss_generation_service.py", "TISSGenerationService — geração de XML a partir do claim"],
        ["Service", "tiss_validation_service.py", "TISSValidationService — validação contra schema XSD"],
        ["DMN", "tiss_xml_generation_validation.dmn", "Regras de validação da geração TISS"],
        ["DMN", "tiss_schema_validation_rules.dmn", "Regras de validação do schema TISS"],
    ],
    col_widths=[3, 7, 8],
)

add_gap_box(doc,
    "Falta processo de emissão de NFS-e municipal para pacientes particulares. "
    "Para o portal de cobrança de exames de imagem, será necessário criar: "
    "(1) Worker de integração com prefeitura (API NFS-e), "
    "(2) DMN de regras fiscais (ISS, CNAE, alíquota), "
    "(3) BPMN de emissão/cancelamento de NFS-e."
)

# ═══════════════════════════════════════════════════════════════════
# 5. CONCILIACAO BANCARIA
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("5. Conciliação Bancária / Reconciliation", level=1)

doc.add_heading("5.1 Processos BPMN", level=2)
add_table(doc,
    ["Processo", "ID", "Descrição"],
    [
        ["SP-RC-010\nPayment Reconciliation",
         "SP_RC_010",
         "Classificar tipo de pagamento → parallel (detectar duplicata + matching automático) → "
         "matched? SIM: ajustes contratuais → calcular pagamento líquido → checar overpayment → "
         "persistir → conciliar mensal. NÃO: flag discrepâncias → revisão manual (timer 24h)."],
        ["SP-RC-012\nCollection Matching",
         "SP_RC_012",
         "Parse arquivo bancário → validar dados → classificar tipo → estratégia de matching "
         "(protocolo / fatura / paciente) → matching automático → detectar duplicata → "
         "converter moeda → ajustes → calcular líquido → gateway (over/under/exato) → "
         "persistir → conciliar (diário/semanal/mensal) → arquivar."],
        ["SP-RC-008\nRevenue Collection",
         "SP_RC_008",
         "Identificar vencidos → priorizar → DMN método → DMN desconto → "
         "receber notificação (timer 7 dias) → matching automático → "
         "conciliar diariamente → atualizar contas a receber."],
    ],
    col_widths=[4, 3, 11],
)

doc.add_heading("5.2 Workers de Conciliação (~23 workers)", level=2)
doc.add_paragraph("Path: revenue_cycle/collection/workers/")
add_table(doc,
    ["Worker", "TOPIC (collection.*)", "Descrição"],
    [
        ["parse_payment_file", "parse_payment_file", "Parse de arquivos bancários CNAB 240/400"],
        ["validate_payment_data", "validate_payment_data", "Validação de integridade dos dados"],
        ["classify_payment_type", "classify_payment_type", "Classifica: full, partial, advance, refund, adjustment"],
        ["auto_matching", "auto_matching", "Matching automático com score de confiança"],
        ["match_by_protocol", "match_by_protocol", "Matching por número de protocolo"],
        ["match_by_invoice", "match_by_invoice", "Matching por número de fatura"],
        ["match_by_patient", "match_by_patient", "Matching por ID do paciente"],
        ["detect_duplicate_payment", "detect_duplicate_payment", "Detecção de pagamentos duplicados"],
        ["apply_contractual_adjustments", "apply_contractual_adjustments", "Aplica ajustes contratuais"],
        ["calculate_net_payment", "calculate_net_payment", "Calcula pagamento líquido"],
        ["flag_discrepancies", "flag_discrepancies", "Sinaliza discrepâncias"],
        ["handle_overpayment", "handle_overpayment", "Processa pagamentos excedentes (gera refundId)"],
        ["handle_underpayment", "handle_underpayment", "Processa pagamentos insuficientes"],
        ["persist_payment", "persist_payment", "Persiste pagamentos reconciliados"],
        ["convert_currency", "convert_currency", "Conversão para BRL"],
        ["partial_allocation", "partial_allocation", "Alocação parcial de pagamentos"],
        ["finalize_allocation", "finalize_allocation", "Finaliza alocação pagamento-claim"],
        ["receive_payment_notification", "receive_payment_notification", "Recebe notificações do banco/gateway"],
        ["escalate_unmatched", "escalate_unmatched", "Escalona pagamentos não conciliados"],
        ["reconcile_daily", "reconcile_daily", "Conciliação diária"],
        ["reconcile_weekly", "reconcile_weekly", "Conciliação semanal"],
        ["reconcile_monthly", "reconcile_monthly", "Conciliação mensal"],
        ["archive_reconciliation", "archive_reconciliation", "Arquiva dados de conciliação"],
    ],
    col_widths=[5, 5.5, 7.5],
)

doc.add_heading("5.3 Bibliotecas de Suporte", level=2)
add_table(doc,
    ["Arquivo", "Descrição"],
    [
        ["cnab_parser.py", "Parser CNAB 240/400 (FEBRABAN). Segmentos T/U/J/A. "
         "Status: LIQUIDACAO_NORMAL, LIQUIDACAO_PARCIAL, ENTRADA_CONFIRMADA, "
         "ENTRADA_REJEITADA, BAIXA, PROTESTADO."],
        ["penalty_calculator.py", "Cálculo de juros e multa (Lei 10.406/2002 — Código Civil). "
         "Juros de mora 1%/mês pro rata die, multa até 2%, correção INPC."],
    ],
    col_widths=[5, 13],
)

doc.add_heading("5.4 Serviços de Reconciliação", level=2)
add_table(doc,
    ["Serviço", "Classe", "Descrição"],
    [
        ["reconciliation_service.py", "ReconciliationService", "Reconciliação cross-source via DMN"],
        ["payment_matching_service.py", "PaymentMatchingService", "Matching por estratégia (invoice/patient/protocol)"],
    ],
    col_widths=[6, 5, 7],
)

# ═══════════════════════════════════════════════════════════════════
# 6. CANCELAMENTO
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("6. Cancelamento / Refund / Estorno", level=1)

doc.add_heading("6.1 O Que Existe", level=2)
doc.add_paragraph(
    "Não existe um BPMN dedicado para cancelamento/estorno/chargeback. "
    "A lógica está distribuída nos seguintes componentes:"
)
add_table(doc,
    ["Componente", "Onde", "O Que Faz"],
    [
        ["handle_overpayment", "SP-RC-010, SP-RC-012",
         "Gera refundId/refundStatus. Error boundary Error_OverpaymentDetected."],
        ["handle_overpayment_worker.py", "collection/workers/",
         "Processa pagamentos excedentes, inicia fluxo de reembolso"],
        ["PaymentType.REFUND", "collection/enums.py",
         "Tipo de pagamento para estorno"],
        ["PaymentType.ADJUSTMENT", "collection/enums.py",
         "Tipo de pagamento para ajuste"],
        ["PaymentStatus.REVERSED", "collection/enums.py",
         "Status para pagamento revertido"],
    ],
    col_widths=[5, 4.5, 8.5],
)

add_gap_box(doc,
    "Falta BPMN dedicado para cancelamento de cobrança, estorno de pagamento e chargeback. "
    "Para o portal, criar: (1) Processo de cancelamento de cobrança (antes do pagamento), "
    "(2) Processo de estorno/reembolso (após pagamento), "
    "(3) Processo de chargeback de cartão de crédito, "
    "(4) Workers de integração com gateway para reversal/refund."
)

# ═══════════════════════════════════════════════════════════════════
# 7. DUNNING
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("7. Cobrança de Inadimplentes / Dunning", level=1)

doc.add_heading("7.1 Processos BPMN", level=2)
add_table(doc,
    ["Processo", "ID", "Descrição"],
    [
        ["SP-RC-008\nRevenue Collection",
         "SP_RC_008",
         "Core: identificar vencidos → priorizar → receber pagamento → matching → conciliar. Timer 7 dias."],
        ["SP-RC-009\nCollection Management",
         "SP_RC_009",
         "Priorização por aging: CRITICAL → jurídico; HIGH → carta + ligação (timer 7d); "
         "MEDIUM → carta; LOW → WhatsApp."],
        ["SP-RC-011\nCollection Analytics",
         "SP_RC_011",
         "Analytics: DSO, taxa de recebimento, variância, tempo do ciclo, performance de operadora, "
         "leakage, dashboard executivo, previsões, alertas."],
        ["SP-RC-013\nCollection Escalation",
         "SP_RC_013",
         "Escalonamento: aging → priorizar → ações por prioridade → sem resposta? → "
         "jurídico ou write-off. Operações: aging report, daily summary, ERP export, BI/DW."],
    ],
    col_widths=[4, 3, 11],
)

doc.add_heading("7.2 Workers Principais (~40 workers)", level=2)
add_table(doc,
    ["Worker", "Descrição"],
    [
        ["identify_overdue", "Identifica contas vencidas por data"],
        ["prioritize_collection", "Prioriza por aging/valor/operadora"],
        ["calculate_aging_bucket", "Calcula buckets (0-30, 31-60... 180+ dias)"],
        ["generate_collection_letter", "Gera cartas em PT (aviso amigável, 2ª notificação, extrajudicial)"],
        ["schedule_collection_call", "Agenda ligações de cobrança"],
        ["send_whatsapp_reminder", "Envia lembretes via WhatsApp"],
        ["escalate_to_legal", "Escalona para departamento jurídico"],
        ["write_off_bad_debt", "Baixa de inadimplência"],
        ["negotiate_payment_plan", "Negocia planos de pagamento"],
        ["apply_penalties", "Aplica multa (2%) + juros (1%/mês) — Lei brasileira"],
        ["predict_collection_date", "Previsão de data de recebimento (ML)"],
        ["identify_slow_payers", "Identifica operadoras lentas"],
        ["calculate_dso / collection_rate / variance", "Métricas de performance"],
        ["detect_revenue_leakage", "Detecta vazamento de receita"],
        ["export_to_erp", "Exporta para ERP (Tasy/MvSoul)"],
        ["update_bi_datawarehouse", "Atualiza BI/DW"],
    ],
    col_widths=[7, 11],
)

doc.add_heading("7.3 Templates de Cartas", level=2)
doc.add_paragraph(
    "collection_letters.py — Templates em português: FIRST_NOTICE (aviso amigável, 15 dias), "
    "SECOND_NOTICE (segunda notificação com juros/multa), notificação extrajudicial."
)

# ═══════════════════════════════════════════════════════════════════
# 8. GLOSA
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("8. Glosa / Denial Management", level=1)

doc.add_heading("8.1 Processo BPMN", level=2)
doc.add_paragraph(
    "SP-RC-007_Denial_Management.bpmn — Identificar glosa → classificar tipo → "
    "analisar razão → DMN predict risk → DMN prevention → check appeal eligibility "
    "(timer 30 dias) → gerar documentação de recurso → submeter → rastrear. "
    "Se não elegível: calcular impacto → escalonar → atualizar pagamento."
)

doc.add_heading("8.2 Workers (10 workers)", level=2)
add_table(doc,
    ["Worker", "TOPIC", "Descrição"],
    [
        ["identify_glosa", "glosa.identify", "Identifica glosas da resposta da operadora"],
        ["classify_glosa_type", "glosa.classify_type", "Classifica: administrativa, técnica, médica"],
        ["analyze_glosa_reason", "glosa.analyze_reason", "Análise de causa raiz"],
        ["check_appeal_eligibility", "glosa.check_appeal_eligibility", "Verifica elegibilidade para recurso"],
        ["generate_appeal_documentation", "glosa.generate_appeal_documentation", "Gera documentação de recurso"],
        ["submit_appeal", "glosa.submit_appeal", "Submete recurso à operadora"],
        ["track_appeal_status", "glosa.track_appeal_status", "Rastreia status do recurso"],
        ["calculate_glosa_impact", "glosa.calculate_impact", "Calcula impacto financeiro"],
        ["escalate_to_supervisor", "glosa.escalate_to_supervisor", "Escalona glosas não resolvidas"],
        ["update_payment", "glosa.update_payment", "Atualiza pagamento após resolução"],
    ],
    col_widths=[5.5, 6, 6.5],
)

doc.add_heading("8.3 DMN de Glosa (~65 arquivos)", level=2)
add_table(doc,
    ["Subdiretório", "Qtd", "Propósito"],
    [
        ["predict/", "12", "Predição de risco de glosa"],
        ["prevent/", "5", "Estratégias de prevenção"],
        ["appeal/", "6", "Regras de elegibilidade de recurso"],
        ["payer/", "10", "Regras por operadora"],
        ["medical/", "7", "Necessidade médica"],
        ["missing/", "7", "Documentação faltante"],
        ["duplicate/", "7", "Cobrança duplicada"],
        ["timing/", "7", "Regras de prazo"],
        ["classification/", "1", "Classificação de tipo de glosa"],
        ["identification/", "1", "Identificação de glosa"],
        ["reason_analysis/", "1", "Análise de razão"],
    ],
    col_widths=[4, 1.5, 12.5],
)

# ═══════════════════════════════════════════════════════════════════
# 9. MAXIMIZACAO
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("9. Maximização e Revenue Optimization", level=1)

doc.add_heading("9.1 SP-RC-010 Maximization", level=2)
doc.add_paragraph(
    "SP-RC-010_Maximization.bpmn — Prever data de recebimento → identificar slow payers → "
    "gateway de risco → DMN estratégia de otimização → atualizar previsões → "
    "analisar oportunidades → parallel (otimizar pricing + otimizar contratos) → "
    "gerar recomendações."
)

doc.add_heading("9.2 Revenue Optimization (Platform Services)", level=2)
doc.add_paragraph(
    "revenue_optimization.bpmn — 14 etapas lineares: identify coding opportunities → "
    "suggest documentation improvements → analyze denial patterns → optimize pricing → "
    "identify contract gaps → recommend bundles → calculate revenue potential → "
    "prioritize high-value cases → detect leakage → optimize resources → "
    "forecast trends → benchmark payers → generate report → track ROI."
)

doc.add_heading("9.3 Workers de Revenue Optimization (14 workers)", level=2)
add_table(doc,
    ["TOPIC", "Descrição"],
    [
        ["platform.analyze_financial_performance", "Análise de performance financeira"],
        ["platform.analyze_denial_patterns", "Análise de padrões de glosa"],
        ["platform.optimize_pricing_strategy", "Otimização de precificação"],
        ["platform.identify_contract_gaps", "Identificação de gaps contratuais"],
        ["platform.recommend_procedure_bundles", "Recomendação de pacotes de procedimentos"],
        ["platform.calculate_revenue_potential", "Cálculo de potencial de receita"],
        ["platform.prioritize_high_value_cases", "Priorização de casos de alto valor"],
        ["platform.detect_revenue_leakage", "Detecção de vazamento de receita"],
        ["platform.forecast_revenue_trends", "Previsão de tendências de receita"],
        ["platform.benchmark_payer_performance", "Benchmarking de operadoras"],
        ["platform.generate_optimization_report", "Geração de relatório"],
        ["platform.track_optimization_roi", "Rastreamento de ROI"],
        ["platform.sync_erp_data", "Sincronização com ERP"],
        ["platform.reconcile_data_sources", "Reconciliação multi-fonte"],
    ],
    col_widths=[8, 10],
)

# ═══════════════════════════════════════════════════════════════════
# 10. FINANCIAL CLEARANCE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Financial Clearance (Patient Access)", level=1)

doc.add_paragraph(
    "SP-PA-004_Financial_Clearance.bpmn — Agendamento → verificar convênio → "
    "DMN classificação de urgência → autorização necessária? SIM: checar autorização → "
    "DMN pré-auth → checar pré-autorização (SLA 48h) → escalonar. "
    "Merge → conceder financial clearance."
)
add_table(doc,
    ["TOPIC", "Descrição"],
    [
        ["patient.verify_insurance", "Verifica dados do convênio"],
        ["patient.check_authorization", "Checa autorização do procedimento"],
        ["scheduling.check_pre_auth", "Verifica pré-autorização"],
        ["patient.escalate_auth_sla", "Escalona SLA de autorização"],
        ["patient.grant_clearance", "Concede clearance financeiro"],
    ],
    col_widths=[6, 12],
)
doc.add_paragraph(
    "DMN de autorização: ~70 arquivos em patient_access/dmn/authorization/ "
    "(preauth, urgency, timing, appeal, coding, documentation, extension, scope, status, track, units, federated)."
)

# ═══════════════════════════════════════════════════════════════════
# 11. PRODUCAO CLINICA
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("11. Produção Clínica e Precificação", level=1)

doc.add_paragraph(
    "SP-RC-004_Clinical_Production.bpmn — Pipeline de captura e precificação de procedimentos."
)
add_table(doc,
    ["Worker", "TOPIC", "Descrição"],
    [
        ["capture_procedure", "revenue_cycle.production.capture_procedure", "Captura procedimentos clínicos"],
        ["enrich_procedure", "revenue_cycle.production.enrich_procedure", "Enriquece com dados clínicos"],
        ["validate_procedure", "revenue_cycle.production.validate_procedure", "Valida códigos de procedimento"],
        ["validate_compatibility", "revenue_cycle.production.validate_compatibility", "Valida compatibilidade"],
        ["calculate_quantity", "revenue_cycle.production.calculate_value", "Calcula quantidades/valores"],
        ["assign_prices", "revenue_cycle.production.assign_prices", "Atribui preços aos procedimentos"],
        ["check_authorization", "revenue_cycle.production.check_authorization", "Checa autorização"],
        ["persist_production", "revenue_cycle.production.record_production", "Persiste registros de produção"],
    ],
    col_widths=[4.5, 8, 5.5],
)

doc.add_heading("Serviços de Produção", level=2)
add_table(doc,
    ["Serviço", "Descrição"],
    [
        ["pricing_assignment_service.py", "Atribuição de preços contratuais"],
        ["procedure_capture_service.py", "Captura de procedimentos"],
        ["procedure_enrichment_service.py", "Enriquecimento de dados clínicos"],
        ["production_persistence_service.py", "Persistência de produção"],
    ],
    col_widths=[7, 11],
)

# ═══════════════════════════════════════════════════════════════════
# 12. RESUMO: EXISTE vs GAPS
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("12. Resumo: O Que Existe vs. Gaps para o Portal", level=1)

doc.add_heading("12.1 O Que JÁ Existe e Pode Ser Reutilizado", level=2)
add_table(doc,
    ["#", "Componente", "Reutilização"],
    [
        ["1", "SP-RC-014 (Self-Service)", "Base do portal — paths de estimativa, pagamento (PAY_NOW/PAY_LATER) e recibo"],
        ["2", "SP-RC-006 (Billing)", "Pipeline completo de faturamento TISS para operadoras"],
        ["3", "SP-RC-010/012 (Reconciliação)", "Matching + conciliação + CNAB parser"],
        ["4", "Enums (PaymentMethod/Status)", "Já define boleto, PIX, cartão, etc."],
        ["5", "Cash Operations DMN", "Regras de desconto, estimativa e copay"],
        ["6", "cnab_parser.py", "Parser CNAB 240/400 pronto para uso"],
        ["7", "penalty_calculator.py", "Cálculo de juros/multa conforme legislação brasileira"],
        ["8", "collection_letters.py", "Templates de cartas de cobrança em PT"],
        ["9", "selfservice_eligibility_rules.dmn", "Elegibilidade de self-service (inclui action=pay)"],
        ["10", "SP-RC-009/013 (Dunning)", "Cobrança automatizada com WhatsApp/carta/ligação"],
    ],
    col_widths=[1, 6.5, 10.5],
)

doc.add_heading("12.2 Gaps Identificados", level=2)
add_table(doc,
    ["#", "Gap", "Prioridade", "Descrição"],
    [
        ["1", "NFS-e Municipal", "ALTA",
         "Emissão de nota fiscal eletrônica para pacientes particulares. "
         "Worker de integração com API da prefeitura, DMN de regras fiscais (ISS, CNAE)."],
        ["2", "Gateway de Pagamento", "ALTA",
         "Workers para: gerar QR Code PIX (API Banco Central), registrar boleto (CNAB), "
         "tokenizar cartão de crédito. Integração com Cielo/Stone/PagSeguro."],
        ["3", "Cancelamento/Estorno", "ALTA",
         "BPMN dedicado: cancelamento de cobrança, estorno de pagamento, "
         "chargeback de cartão, reversal via gateway."],
        ["4", "Precificação de Imagem", "MÉDIA",
         "DMN específico para precificação de exames de imagem "
         "(ressonância, tomografia, raio-x, ultrassom) com copay por convênio."],
        ["5", "Parcelamento Avançado", "MÉDIA",
         "Processo de parcelamento no cartão de crédito, "
         "com juros diferenciados e controle de parcelas."],
        ["6", "Notificação de Cobrança", "MÉDIA",
         "Worker de notificação push/SMS/email para cobrança automática de exames, "
         "com templates específicos para exames de imagem."],
    ],
    col_widths=[1, 4.5, 3, 9.5],
)

# ═══════════════════════════════════════════════════════════════════
# 13. RECOMENDACOES
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("13. Recomendações de Implementação", level=1)

doc.add_heading("13.1 Fase 1 — Portal Básico (MVP)", level=2)
recommendations_phase1 = [
    "Estender SP-RC-014 (Self-Service) com path específico para exames de imagem",
    "Criar workers de gateway de pagamento (PIX QR Code + Boleto)",
    "Integrar com SP-RC-006 (Billing) para geração automática de cobranças",
    "Reutilizar cash_operations DMN para estimativas e descontos",
    "Implementar webhook de confirmação de pagamento (PIX/Boleto)",
]
for i, rec in enumerate(recommendations_phase1, 1):
    doc.add_paragraph(f"{i}. {rec}")

doc.add_heading("13.2 Fase 2 — Compliance e Automação", level=2)
recommendations_phase2 = [
    "Criar BPMN de emissão de NFS-e com integração municipal",
    "Implementar processo de cancelamento/estorno dedicado",
    "Adicionar parcelamento no cartão com regras DMN de juros",
    "Integrar reconciliação automática (SP-RC-012) com arquivo de retorno do gateway",
    "Criar DMN específico de precificação para exames de imagem por convênio/particular",
]
for i, rec in enumerate(recommendations_phase2, 1):
    doc.add_paragraph(f"{i}. {rec}")

doc.add_heading("13.3 Fase 3 — Otimização", level=2)
recommendations_phase3 = [
    "Ativar dunning automatizado (SP-RC-009/013) para cobranças vencidas do portal",
    "Integrar com analytics (SP-RC-011) para dashboards de performance do portal",
    "Implementar previsão ML de inadimplência (predict_collection_date_worker)",
    "Conectar com revenue optimization para identificar oportunidades de pricing",
]
for i, rec in enumerate(recommendations_phase3, 1):
    doc.add_paragraph(f"{i}. {rec}")

# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
output_path = r"c:\BPMN\Healthcare-Orchest-CIB7\docs\data-mapping\Mapeamento_Fluxos_Financeiros_MAEZO.docx"
doc.save(output_path)
print(f"DOCX gerado: {output_path}")
